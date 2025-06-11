"""
增强知识库管理模型 - AI增强版
Enhanced Knowledge Base Management with AI Integration
"""

import os
import json
import hashlib
import numpy as np
from datetime import datetime
from typing import List, Dict, Optional, Tuple
import logging
from sklearn.feature_extraction.text import TfidfVectorizer
from sklearn.metrics.pairwise import cosine_similarity
import jieba
import re

# 文档内容提取库
import openpyxl
import pdfplumber
import docx
from werkzeug.utils import secure_filename

logger = logging.getLogger(__name__)

class EnhancedDocument:
    """增强文档模型"""
    
    def __init__(self, doc_id: int, filename: str, content: str, doc_type: str, 
                 category: str = 'general', title: str = None, description: str = None,
                 upload_time: str = None):
        self.doc_id = doc_id
        self.filename = filename
        self.content = content
        self.doc_type = doc_type
        self.category = category
        self.title = title or filename
        self.description = description or ""
        self.upload_time = upload_time or datetime.now().isoformat()
        
        # AI增强字段
        self.keywords = self._extract_keywords()
        self.summary = self._generate_summary()
        self.content_vector = None  # 将由向量化器设置
        
    def _extract_keywords(self) -> List[str]:
        """提取关键词"""
        if not self.content:
            return []
        
        # 中文分词
        words = jieba.cut(self.content)
        
        # 过滤停用词和短词
        stop_words = {'的', '了', '在', '是', '我', '有', '和', '就', '不', '人', '都', '一', '一个', '上', '也', '很', '到', '说', '要', '去', '你', '会', '着', '没有', '看', '好', '自己', '这'}
        keywords = []
        
        for word in words:
            word = word.strip()
            if (len(word) >= 2 and 
                word not in stop_words and 
                not word.isdigit() and
                re.match(r'^[\u4e00-\u9fa5a-zA-Z0-9\-_]+$', word)):
                keywords.append(word)
        
        # 返回频率最高的关键词
        from collections import Counter
        word_freq = Counter(keywords)
        return [word for word, freq in word_freq.most_common(20)]
    
    def _generate_summary(self, max_length: int = 200) -> str:
        """生成智能摘要"""
        if not self.content:
            return ""
        
        # 简单的摘要生成：提取包含关键信息的句子
        sentences = re.split(r'[。！？\n]', self.content)
        
        # 优先选择包含技术关键词的句子
        tech_keywords = ['PXI', 'LabVIEW', 'C#', '数据采集', '测控', '简仪', 'MISD', '模块', '系统', '平台']
        important_sentences = []
        
        for sentence in sentences:
            sentence = sentence.strip()
            if len(sentence) > 10:
                # 计算句子重要性得分
                score = sum(1 for keyword in tech_keywords if keyword in sentence)
                if score > 0:
                    important_sentences.append((sentence, score))
        
        # 按重要性排序并选择前几句
        important_sentences.sort(key=lambda x: x[1], reverse=True)
        
        summary_parts = []
        current_length = 0
        
        for sentence, score in important_sentences:
            if current_length + len(sentence) <= max_length:
                summary_parts.append(sentence)
                current_length += len(sentence)
            else:
                break
        
        if not summary_parts and sentences:
            # 如果没有找到重要句子，取开头部分
            summary_parts = [sentences[0][:max_length]]
        
        return '。'.join(summary_parts) + ('。' if summary_parts else '')
    
    def to_dict(self) -> Dict:
        """转换为字典"""
        return {
            'doc_id': self.doc_id,
            'filename': self.filename,
            'content': self.content,
            'doc_type': self.doc_type,
            'category': self.category,
            'title': self.title,
            'description': self.description,
            'upload_time': self.upload_time,
            'keywords': self.keywords,
            'summary': self.summary
        }

class EnhancedContentExtractor:
    """增强内容提取器"""
    
    @staticmethod
    def extract_excel_content(file_path: str) -> str:
        """提取Excel文件内容"""
        try:
            workbook = openpyxl.load_workbook(file_path, data_only=True)
            content_parts = []
            
            for sheet_name in workbook.sheetnames:
                sheet = workbook[sheet_name]
                content_parts.append(f"工作表: {sheet_name}")
                
                for row in sheet.iter_rows(values_only=True):
                    row_data = [str(cell) if cell is not None else "" for cell in row]
                    row_text = " | ".join(filter(None, row_data))
                    if row_text.strip():
                        content_parts.append(row_text)
            
            return "\n".join(content_parts)
            
        except Exception as e:
            logger.error(f"Excel内容提取失败: {e}")
            return ""
    
    @staticmethod
    def extract_pdf_content(file_path: str) -> str:
        """提取PDF文件内容（使用pdfplumber）"""
        try:
            content_parts = []
            
            with pdfplumber.open(file_path) as pdf:
                for page_num, page in enumerate(pdf.pages):
                    text = page.extract_text()
                    if text:
                        content_parts.append(f"第{page_num + 1}页:\n{text}")
                    
                    # 尝试提取表格
                    tables = page.extract_tables()
                    for table_num, table in enumerate(tables):
                        if table:
                            table_text = f"表格{table_num + 1}:\n"
                            for row in table:
                                if row:
                                    row_text = " | ".join(str(cell) if cell else "" for cell in row)
                                    table_text += row_text + "\n"
                            content_parts.append(table_text)
            
            return "\n\n".join(content_parts)
            
        except Exception as e:
            logger.error(f"PDF内容提取失败: {e}")
            return ""
    
    @staticmethod
    def extract_word_content(file_path: str) -> str:
        """提取Word文档内容"""
        try:
            doc = docx.Document(file_path)
            content_parts = []
            
            # 提取段落
            for paragraph in doc.paragraphs:
                if paragraph.text.strip():
                    content_parts.append(paragraph.text)
            
            # 提取表格
            for table_num, table in enumerate(doc.tables):
                table_text = f"表格{table_num + 1}:\n"
                for row in table.rows:
                    row_text = " | ".join(cell.text.strip() for cell in row.cells)
                    if row_text.strip():
                        table_text += row_text + "\n"
                content_parts.append(table_text)
            
            return "\n".join(content_parts)
            
        except Exception as e:
            logger.error(f"Word内容提取失败: {e}")
            return ""
    
    @staticmethod
    def extract_text_content(file_path: str) -> str:
        """提取文本文件内容"""
        try:
            encodings = ['utf-8', 'gbk', 'gb2312', 'latin-1']
            
            for encoding in encodings:
                try:
                    with open(file_path, 'r', encoding=encoding) as f:
                        return f.read()
                except UnicodeDecodeError:
                    continue
            
            logger.warning(f"无法解码文本文件: {file_path}")
            return ""
            
        except Exception as e:
            logger.error(f"文本内容提取失败: {e}")
            return ""

class VectorSearchEngine:
    """向量搜索引擎"""
    
    def __init__(self):
        self.tfidf_vectorizer = TfidfVectorizer(
            max_features=5000,
            stop_words=None,  # 我们自己处理中文停用词
            ngram_range=(1, 2),
            min_df=1,
            max_df=0.95
        )
        self.document_vectors = None
        self.documents = []
        self.is_fitted = False
    
    def preprocess_text(self, text: str) -> str:
        """预处理文本"""
        if not text:
            return ""
        
        # 中文分词
        words = jieba.cut(text)
        
        # 过滤和清理
        processed_words = []
        stop_words = {'的', '了', '在', '是', '我', '有', '和', '就', '不', '人', '都', '一', '一个', '上', '也', '很', '到', '说', '要', '去', '你', '会', '着', '没有', '看', '好', '自己', '这'}
        
        for word in words:
            word = word.strip().lower()
            if (len(word) >= 2 and 
                word not in stop_words and 
                not word.isdigit()):
                processed_words.append(word)
        
        return " ".join(processed_words)
    
    def fit_documents(self, documents: List[EnhancedDocument]):
        """训练向量化器"""
        self.documents = documents
        
        if not documents:
            logger.warning("没有文档用于训练向量化器")
            return
        
        # 预处理所有文档内容
        processed_texts = []
        for doc in documents:
            # 组合标题、描述和内容
            combined_text = f"{doc.title} {doc.description} {doc.content}"
            processed_text = self.preprocess_text(combined_text)
            processed_texts.append(processed_text)
        
        try:
            # 训练TF-IDF向量化器
            self.document_vectors = self.tfidf_vectorizer.fit_transform(processed_texts)
            self.is_fitted = True
            logger.info(f"向量搜索引擎训练完成，文档数量: {len(documents)}")
            
        except Exception as e:
            logger.error(f"向量化器训练失败: {e}")
            self.is_fitted = False
    
    def search(self, query: str, top_k: int = 5) -> List[Tuple[EnhancedDocument, float]]:
        """向量搜索"""
        if not self.is_fitted or not self.documents:
            return []
        
        try:
            # 预处理查询
            processed_query = self.preprocess_text(query)
            if not processed_query:
                return []
            
            # 向量化查询
            query_vector = self.tfidf_vectorizer.transform([processed_query])
            
            # 计算相似度
            similarities = cosine_similarity(query_vector, self.document_vectors).flatten()
            
            # 获取最相似的文档
            top_indices = similarities.argsort()[-top_k:][::-1]
            
            results = []
            for idx in top_indices:
                if similarities[idx] > 0.01:  # 最小相似度阈值
                    results.append((self.documents[idx], float(similarities[idx])))
            
            return results
            
        except Exception as e:
            logger.error(f"向量搜索失败: {e}")
            return []

class EnhancedKnowledgeBase:
    """增强知识库管理类"""
    
    def __init__(self):
        from models.database import document_manager
        self.document_manager = document_manager
        self.content_extractor = EnhancedContentExtractor()
        self.vector_engine = VectorSearchEngine()
        self.documents_cache = {}
        self._load_and_index_documents()
    
    def _load_and_index_documents(self):
        """加载并索引所有文档"""
        try:
            # 从数据库获取所有文档
            db_docs = self.document_manager.get_all_documents(per_page=1000)
            enhanced_docs = []
            
            for db_doc in db_docs['documents']:
                # 重新提取内容（使用增强提取器）
                content = self._extract_enhanced_content(db_doc)
                
                if content:  # 只处理有内容的文档
                    enhanced_doc = EnhancedDocument(
                        doc_id=db_doc['id'],
                        filename=db_doc['original_filename'],
                        content=content,
                        doc_type=db_doc['file_type'],
                        category=db_doc['category'],
                        title=db_doc['title'],
                        description=db_doc['description'],
                        upload_time=db_doc['upload_time']
                    )
                    enhanced_docs.append(enhanced_doc)
                    self.documents_cache[db_doc['id']] = enhanced_doc
            
            # 训练向量搜索引擎
            if enhanced_docs:
                self.vector_engine.fit_documents(enhanced_docs)
                logger.info(f"知识库索引完成，有效文档数量: {len(enhanced_docs)}")
            else:
                logger.warning("没有找到有效的文档内容")
                
        except Exception as e:
            logger.error(f"知识库索引失败: {e}")
    
    def _extract_enhanced_content(self, db_doc: Dict) -> str:
        """使用增强提取器重新提取文档内容"""
        try:
            # 构建文件路径
            file_path = os.path.join(
                os.path.dirname(__file__), '..', 'data', 'uploads', db_doc['filename']
            )
            
            if not os.path.exists(file_path):
                logger.warning(f"文件不存在: {file_path}")
                return db_doc.get('content_text', '')
            
            file_type = db_doc['file_type']
            
            if file_type == 'excel':
                return self.content_extractor.extract_excel_content(file_path)
            elif file_type == 'pdf':
                return self.content_extractor.extract_pdf_content(file_path)
            elif file_type == 'word':
                return self.content_extractor.extract_word_content(file_path)
            elif file_type in ['text', 'markdown']:
                return self.content_extractor.extract_text_content(file_path)
            else:
                return db_doc.get('content_text', '')
                
        except Exception as e:
            logger.error(f"增强内容提取失败: {e}")
            return db_doc.get('content_text', '')
    
    def search_documents(self, query: str, limit: int = 5) -> List[Dict]:
        """增强文档搜索"""
        try:
            # 使用向量搜索
            vector_results = self.vector_engine.search(query, top_k=limit)
            
            # 转换为标准格式
            results = []
            for doc, score in vector_results:
                results.append({
                    'id': doc.doc_id,
                    'title': doc.title,
                    'original_filename': doc.filename,
                    'file_type': doc.doc_type,
                    'category': doc.category,
                    'content_summary': doc.summary,
                    'upload_time': doc.upload_time,
                    'relevance_score': score,
                    'keywords': doc.keywords
                })
            
            return results
            
        except Exception as e:
            logger.error(f"增强搜索失败: {e}")
            return []
    
    def get_relevant_content(self, question: str, max_docs: int = 3) -> str:
        """获取与问题相关的文档内容（RAG增强）"""
        try:
            # 使用向量搜索找到最相关的文档
            vector_results = self.vector_engine.search(question, top_k=max_docs)
            
            if not vector_results:
                logger.warning(f"向量搜索没有找到相关文档，问题: {question}")
                return ""
            
            content_parts = []
            
            for doc, score in vector_results:
                logger.info(f"处理文档: {doc.title}, 相似度: {score:.3f}")
                
                if score < 0.01:  # 降低相似度阈值
                    logger.info(f"文档相似度太低，跳过: {score:.3f}")
                    continue
                
                # 提取最相关的段落
                relevant_paragraphs = self._extract_relevant_paragraphs(
                    doc.content, question, max_paragraphs=3
                )
                
                if relevant_paragraphs:
                    content_part = f"【{doc.title}】(相关度: {score:.2f})\n"
                    content_part += "\n".join(relevant_paragraphs)
                    content_parts.append(content_part)
                    logger.info(f"添加了 {len(relevant_paragraphs)} 个相关段落")
                else:
                    # 如果没有找到相关段落，使用文档摘要
                    if doc.summary:
                        content_part = f"【{doc.title}】(相关度: {score:.2f})\n{doc.summary}"
                        content_parts.append(content_part)
                        logger.info(f"使用文档摘要作为相关内容")
                    else:
                        # 使用文档内容的前500个字符
                        preview = doc.content[:500] if doc.content else ""
                        if preview:
                            content_part = f"【{doc.title}】(相关度: {score:.2f})\n{preview}..."
                            content_parts.append(content_part)
                            logger.info(f"使用文档内容预览")
            
            result = "\n\n".join(content_parts)
            logger.info(f"最终相关内容长度: {len(result)}")
            return result
            
        except Exception as e:
            logger.error(f"相关内容提取失败: {e}")
            return ""
    
    def _extract_relevant_paragraphs(self, content: str, question: str, max_paragraphs: int = 3) -> List[str]:
        """从文档中提取最相关的段落"""
        if not content:
            return []
        
        # 分割段落
        paragraphs = [p.strip() for p in content.split('\n') if len(p.strip()) > 20]
        
        if not paragraphs:
            return []
        
        # 提取问题关键词
        question_words = set(jieba.cut(question.lower()))
        question_words = {w for w in question_words if len(w) >= 2}
        
        # 计算每个段落的相关性得分
        paragraph_scores = []
        
        for paragraph in paragraphs:
            paragraph_words = set(jieba.cut(paragraph.lower()))
            
            # 计算关键词匹配得分
            match_score = len(question_words & paragraph_words) / max(len(question_words), 1)
            
            # 技术关键词加权
            tech_keywords = {'pxi', 'labview', 'c#', '数据采集', '测控', '简仪', 'misd', '模块'}
            tech_score = sum(1 for keyword in tech_keywords if keyword in paragraph.lower())
            
            total_score = match_score + tech_score * 0.1
            paragraph_scores.append((paragraph, total_score))
        
        # 按得分排序并返回最相关的段落
        paragraph_scores.sort(key=lambda x: x[1], reverse=True)
        
        return [p[0] for p in paragraph_scores[:max_paragraphs] if p[1] > 0]
    
    def refresh_index(self):
        """刷新索引"""
        self.documents_cache.clear()
        self._load_and_index_documents()
    
    def get_document_statistics(self) -> Dict:
        """获取文档统计信息"""
        try:
            total_docs = len(self.documents_cache)
            
            # 按类型统计
            type_stats = {}
            category_stats = {}
            
            for doc in self.documents_cache.values():
                type_stats[doc.doc_type] = type_stats.get(doc.doc_type, 0) + 1
                category_stats[doc.category] = category_stats.get(doc.category, 0) + 1
            
            return {
                'total_documents': total_docs,
                'indexed_documents': total_docs,
                'by_type': type_stats,
                'by_category': category_stats,
                'vector_engine_status': 'active' if self.vector_engine.is_fitted else 'inactive'
            }
            
        except Exception as e:
            logger.error(f"获取统计信息失败: {e}")
            return {
                'total_documents': 0,
                'indexed_documents': 0,
                'by_type': {},
                'by_category': {},
                'vector_engine_status': 'error'
            }

# 全局增强知识库实例
enhanced_knowledge_base = EnhancedKnowledgeBase()
