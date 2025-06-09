"""
知识库管理模型
Knowledge Base Management Models
"""

import os
import json
import hashlib
from datetime import datetime
from typing import List, Dict, Optional
import PyPDF2
import docx
from werkzeug.utils import secure_filename

class Document:
    """文档模型"""
    
    def __init__(self, filename: str, content: str, doc_type: str, upload_time: str = None):
        self.filename = filename
        self.content = content
        self.doc_type = doc_type
        self.upload_time = upload_time or datetime.now().isoformat()
        self.doc_id = self._generate_id()
        
    def _generate_id(self) -> str:
        """生成文档ID"""
        content_hash = hashlib.md5(self.content.encode('utf-8')).hexdigest()
        return f"{self.filename}_{content_hash[:8]}"
    
    def to_dict(self) -> Dict:
        """转换为字典"""
        return {
            'doc_id': self.doc_id,
            'filename': self.filename,
            'content': self.content,
            'doc_type': self.doc_type,
            'upload_time': self.upload_time
        }
    
    @classmethod
    def from_dict(cls, data: Dict) -> 'Document':
        """从字典创建文档对象"""
        doc = cls(
            filename=data['filename'],
            content=data['content'],
            doc_type=data['doc_type'],
            upload_time=data['upload_time']
        )
        doc.doc_id = data['doc_id']
        return doc

class KnowledgeBase:
    """知识库管理类"""
    
    def __init__(self, storage_path: str = "src/uploads"):
        self.storage_path = storage_path
        self.documents_file = os.path.join(storage_path, "documents.json")
        self.documents: List[Document] = []
        self._ensure_storage_path()
        self._load_documents()
    
    def _ensure_storage_path(self):
        """确保存储路径存在"""
        if not os.path.exists(self.storage_path):
            os.makedirs(self.storage_path)
    
    def _load_documents(self):
        """加载已存储的文档"""
        if os.path.exists(self.documents_file):
            try:
                with open(self.documents_file, 'r', encoding='utf-8') as f:
                    data = json.load(f)
                    self.documents = [Document.from_dict(doc_data) for doc_data in data]
            except Exception as e:
                print(f"加载文档失败: {e}")
                self.documents = []
    
    def _save_documents(self):
        """保存文档到文件"""
        try:
            with open(self.documents_file, 'w', encoding='utf-8') as f:
                json.dump([doc.to_dict() for doc in self.documents], f, ensure_ascii=False, indent=2)
        except Exception as e:
            print(f"保存文档失败: {e}")
    
    def extract_text_from_pdf(self, file_path: str) -> str:
        """从PDF文件提取文本"""
        try:
            with open(file_path, 'rb') as file:
                pdf_reader = PyPDF2.PdfReader(file)
                text = ""
                for page in pdf_reader.pages:
                    text += page.extract_text() + "\n"
                return text.strip()
        except Exception as e:
            print(f"PDF文本提取失败: {e}")
            return ""
    
    def extract_text_from_docx(self, file_path: str) -> str:
        """从Word文档提取文本"""
        try:
            doc = docx.Document(file_path)
            text = ""
            for paragraph in doc.paragraphs:
                text += paragraph.text + "\n"
            return text.strip()
        except Exception as e:
            print(f"Word文档文本提取失败: {e}")
            return ""
    
    def extract_text_from_txt(self, file_path: str) -> str:
        """从文本文件提取内容"""
        try:
            with open(file_path, 'r', encoding='utf-8') as file:
                return file.read().strip()
        except UnicodeDecodeError:
            try:
                with open(file_path, 'r', encoding='gbk') as file:
                    return file.read().strip()
            except Exception as e:
                print(f"文本文件读取失败: {e}")
                return ""
        except Exception as e:
            print(f"文本文件读取失败: {e}")
            return ""
    
    def upload_document(self, file, category: str = "general") -> Dict:
        """上传文档"""
        try:
            # 安全的文件名
            filename = secure_filename(file.filename)
            if not filename:
                return {"success": False, "error": "无效的文件名"}
            
            # 检查文件类型
            allowed_extensions = {'.pdf', '.docx', '.doc', '.txt', '.md'}
            file_ext = os.path.splitext(filename)[1].lower()
            if file_ext not in allowed_extensions:
                return {"success": False, "error": "不支持的文件类型"}
            
            # 保存文件
            file_path = os.path.join(self.storage_path, filename)
            file.save(file_path)
            
            # 提取文本内容
            content = ""
            if file_ext == '.pdf':
                content = self.extract_text_from_pdf(file_path)
            elif file_ext in ['.docx', '.doc']:
                content = self.extract_text_from_docx(file_path)
            elif file_ext in ['.txt', '.md']:
                content = self.extract_text_from_txt(file_path)
            
            if not content:
                os.remove(file_path)  # 删除无法提取内容的文件
                return {"success": False, "error": "无法提取文件内容"}
            
            # 创建文档对象
            document = Document(
                filename=filename,
                content=content,
                doc_type=category
            )
            
            # 检查是否已存在相同文档
            existing_doc = self.find_document_by_id(document.doc_id)
            if existing_doc:
                os.remove(file_path)  # 删除重复文件
                return {"success": False, "error": "文档已存在"}
            
            # 添加到知识库
            self.documents.append(document)
            self._save_documents()
            
            return {
                "success": True,
                "document": document.to_dict(),
                "message": f"文档 {filename} 上传成功"
            }
            
        except Exception as e:
            return {"success": False, "error": f"上传失败: {str(e)}"}
    
    def find_document_by_id(self, doc_id: str) -> Optional[Document]:
        """根据ID查找文档"""
        for doc in self.documents:
            if doc.doc_id == doc_id:
                return doc
        return None
    
    def search_documents(self, query: str, limit: int = 5) -> List[Document]:
        """搜索文档"""
        query_lower = query.lower()
        results = []
        
        for doc in self.documents:
            # 简单的关键词匹配
            if (query_lower in doc.filename.lower() or 
                query_lower in doc.content.lower()):
                results.append(doc)
        
        # 按相关性排序（简单实现：按匹配次数）
        def relevance_score(doc):
            content_lower = doc.content.lower()
            filename_lower = doc.filename.lower()
            score = content_lower.count(query_lower) * 1 + filename_lower.count(query_lower) * 2
            return score
        
        results.sort(key=relevance_score, reverse=True)
        return results[:limit]
    
    def get_all_documents(self) -> List[Document]:
        """获取所有文档"""
        return self.documents
    
    def get_documents_by_type(self, doc_type: str) -> List[Document]:
        """根据类型获取文档"""
        return [doc for doc in self.documents if doc.doc_type == doc_type]
    
    def delete_document(self, doc_id: str) -> Dict:
        """删除文档"""
        try:
            doc = self.find_document_by_id(doc_id)
            if not doc:
                return {"success": False, "error": "文档不存在"}
            
            # 删除文件
            file_path = os.path.join(self.storage_path, doc.filename)
            if os.path.exists(file_path):
                os.remove(file_path)
            
            # 从列表中移除
            self.documents = [d for d in self.documents if d.doc_id != doc_id]
            self._save_documents()
            
            return {"success": True, "message": "文档删除成功"}
            
        except Exception as e:
            return {"success": False, "error": f"删除失败: {str(e)}"}
    
    def get_relevant_content(self, question: str, max_docs: int = 3) -> str:
        """获取与问题相关的文档内容"""
        relevant_docs = self.search_documents(question, limit=max_docs)
        
        if not relevant_docs:
            return ""
        
        content_parts = []
        for doc in relevant_docs:
            # 提取相关段落（简单实现）
            paragraphs = doc.content.split('\n')
            relevant_paragraphs = []
            
            query_words = question.lower().split()
            for paragraph in paragraphs:
                if any(word in paragraph.lower() for word in query_words):
                    relevant_paragraphs.append(paragraph.strip())
            
            if relevant_paragraphs:
                content_parts.append(f"【{doc.filename}】\n" + "\n".join(relevant_paragraphs[:3]))
            else:
                # 如果没有找到相关段落，取前几段
                content_parts.append(f"【{doc.filename}】\n" + "\n".join(paragraphs[:2]))
        
        return "\n\n".join(content_parts)

# 全局知识库实例
knowledge_base = KnowledgeBase()
